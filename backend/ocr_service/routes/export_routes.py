from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import io

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    pass

router = APIRouter()

class DocumentMeta(BaseModel):
    type: Optional[str] = None
    module: Optional[str] = None
    isRelated: Optional[bool] = None
    isStructured: Optional[bool] = None
    tableDetected: Optional[bool] = None
    confidence: Optional[float] = None
    details: Optional[Dict[str, Any]] = None

class TableData(BaseModel):
    title: Optional[str] = None
    type: Optional[str] = None
    headers: Optional[List[str]] = None
    columns: Optional[List[str]] = None
    rows: Optional[List[Any]] = None

class OcrExportRequest(BaseModel):
    success: Optional[bool] = None
    document: Optional[DocumentMeta] = None
    sections: Optional[List[TableData]] = None
    tables: Optional[List[TableData]] = None
    rawText: Optional[str] = None
    vendor: Optional[Dict[str, Any]] = None
    invoice: Optional[Dict[str, Any]] = None
    totals: Optional[Dict[str, Any]] = None


def _resolve_sections(data: OcrExportRequest) -> List[TableData]:
    """Return sections[], falling back to tables[] for backward compat."""
    return data.sections or data.tables or []


def _get_clean_cols(table_data: TableData) -> List[str]:
    """Get columns."""
    return table_data.headers or table_data.columns or []


# ─────────────────────────────────────────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/docx")
async def export_docx(data: OcrExportRequest, filename: Optional[str] = None):
    try:
        doc = Document()
    except NameError:
        raise HTTPException(status_code=500, detail="python-docx is not installed")

    doc_type = (data.document.type if data.document and data.document.type else "Document")
    sections = _resolve_sections(data)

    # ── Document title ───────────────────────
    title_para = doc.add_heading(doc_type, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Tables ────────────────────────────────────────────────────────────────
    for section in sections:
        cols = _get_clean_cols(section)
        rows = section.rows or []
        if not cols:
            continue

        # Section heading
        if section.title and section.title not in ("Extracted Data", "Extracted Text"):
            doc.add_heading(section.title, level=1)

        # Build table
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(cols):
            hdr_cells[i].text = col_name
            # Bold the header
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True

        # Data rows
        for row in rows:
            row_cells = table.add_row().cells
            for i, col_name in enumerate(cols):
                val = row[i] if isinstance(row, list) and i < len(row) else (row.get(col_name, "") if isinstance(row, dict) else "")
                row_cells[i].text = str(val or "")

        doc.add_paragraph("")  # spacing after table

    # ── Totals block (removed per table-first design) ────────────────────────────────

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    out_filename = filename or f"{doc_type.replace(' ', '_')}_extracted.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={out_filename}"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# TXT EXPORT
# ─────────────────────────────────────────────────────────────────────────────

@router.post("/txt")
async def export_txt(data: OcrExportRequest, filename: Optional[str] = None):
    text_content = data.rawText or "No text extracted."
    buffer = io.BytesIO(text_content.encode("utf-8"))
    out_filename = filename or "extracted.txt"
    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={out_filename}"},
    )


# ─────────────────────────────────────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────────────────────────────────────

try:
    from reportlab.lib import colors  # type: ignore
    from reportlab.lib.pagesizes import letter, landscape  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.lib.enums import TA_CENTER, TA_LEFT  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
except ImportError:
    pass


@router.post("/pdf")
async def export_pdf(data: OcrExportRequest, filename: Optional[str] = None):
    try:
        styles = getSampleStyleSheet()
    except NameError:
        raise HTTPException(status_code=500, detail="reportlab is not installed")

    doc_type = (data.document.type if data.document and data.document.type else "Document")
    sections = _resolve_sections(data)

    buffer = io.BytesIO()

    # Use landscape for tables with many columns, portrait otherwise
    max_cols = max((len(_get_clean_cols(s)) for s in sections), default=0)
    pagesize = landscape(letter) if max_cols > 5 else letter

    pdf_doc = SimpleDocTemplate(
        buffer, pagesize=pagesize,
        rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=24
    )
    elements = []

    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Heading1"],
        alignment=TA_CENTER, fontSize=16, spaceAfter=10
    )
    section_style = ParagraphStyle(
        "SectionHead", parent=styles["Heading2"],
        fontSize=12, spaceBefore=14, spaceAfter=6
    )
    detail_style = ParagraphStyle(
        "Detail", parent=styles["Normal"],
        fontSize=9, spaceAfter=2
    )

    # ── Document title ─────────────────────────────────────────────────────
    elements.append(Paragraph(doc_type, title_style))

    # ── Tables ────────────────────────────────────────────────────────────
    for section in sections:
        cols = _get_clean_cols(section)
        rows = section.rows or []
        if not cols:
            continue

        if section.title and section.title not in ("Extracted Data", "Extracted Text"):
            elements.append(Paragraph(section.title, section_style))

        # Build matrix: header + data rows
        matrix = [cols]
        for row in rows:
            row_data = []
            for i, col in enumerate(cols):
                val = row[i] if isinstance(row, list) and i < len(row) else (row.get(col, "") if isinstance(row, dict) else "")
                row_data.append(str(val or ""))
            matrix.append(row_data)

        # Auto-size columns proportionally
        page_w = pagesize[0] - 72  # full width minus margins
        col_w = page_w / len(cols)
        col_widths = [col_w] * len(cols)

        t = Table(matrix, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            # Header
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f1f5f9")),
            ("TEXTCOLOR",  (0, 0), (-1, 0), colors.HexColor("#374151")),
            ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE",   (0, 0), (-1, 0), 9),
            ("ALIGN",      (0, 0), (-1, 0), "LEFT"),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ("TOPPADDING",    (0, 0), (-1, 0), 8),
            # Data rows
            ("BACKGROUND", (0, 1), (-1, -1), colors.white),
            ("TEXTCOLOR",  (0, 1), (-1, -1), colors.HexColor("#111827")),
            ("FONTNAME",   (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",   (0, 1), (-1, -1), 9),
            ("ALIGN",      (0, 1), (-1, -1), "LEFT"),
            ("TOPPADDING",    (0, 1), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
            # Grid
            ("GRID",    (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
            ("VALIGN",  (0, 0), (-1, -1), "MIDDLE"),
            ("WORDWRAP", (0, 0), (-1, -1), True),
            # Alternating row shading
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f9fafb")]),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 16))

    # ── Totals ────────────────────────────────────────────────────────────

    pdf_doc.build(elements)
    buffer.seek(0)

    out_filename = filename or f"{doc_type.replace(' ', '_')}_extracted.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={out_filename}"},
    )
