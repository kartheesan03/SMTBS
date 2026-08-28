import re
import itertools
from typing import List, Dict, Any, Tuple, Optional
import logging

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# UTILITY
# ---------------------------------------------------------------------------

def parse_number(text: str) -> Optional[float]:
    if not text:
        return None
    clean = re.sub(r'[^\d\.\,\-]', '', text)
    if not clean:
        return None
    if clean.count(',') > 0 and clean.count('.') == 1:
        clean = clean.replace(',', '')
    elif clean.count(',') == 1 and clean.count('.') == 0:
        parts = clean.split(',')
        if len(parts[1]) == 2:
            clean = clean.replace(',', '.')
        else:
            clean = clean.replace(',', '')
    else:
        clean = clean.replace(',', '')
    try:
        return float(clean)
    except Exception:
        return None

# ---------------------------------------------------------------------------
# GEOMETRY & GROUPING
# ---------------------------------------------------------------------------

def merge_boxes(elements: List[Dict]) -> List[Dict]:
    """Merge horizontally adjacent fragmented bounding boxes."""
    # First sort purely by Y
    elements = sorted(elements, key=lambda e: (e['y0'] + e['y1']) / 2)
    merged = []
    
    for el in elements:
        if not merged:
            merged.append(el)
            continue
            
        last = merged[-1]
        
        # Are they on the same line?
        y_mid = (el['y0'] + el['y1']) / 2
        last_y_mid = (last['y0'] + last['y1']) / 2
        height = last['y1'] - last['y0']
        
        # Roughly same line
        if abs(y_mid - last_y_mid) < height * 0.5:
            # Check horizontal gap
            gap = el['x0'] - last['x1']
            if 0 <= gap < height * 0.8:
                # Merge!
                last['text'] = last['text'] + " " + el['text']
                last['x1'] = max(last['x1'], el['x1'])
                last['y0'] = min(last['y0'], el['y0'])
                last['y1'] = max(last['y1'], el['y1'])
                last['confidence'] = (last['confidence'] + el['confidence']) / 2
                continue
        
        merged.append(el)
        
    return merged

def group_into_lines(elements: List[Dict], y_tolerance: int = 15) -> List[List[Dict]]:
    elements = sorted(elements, key=lambda e: ((e['y0'] + e['y1']) / 2, e['x0']))
    lines: List[List[Dict]] = []
    current_line: List[Dict] = []
    current_y: Optional[float] = None

    for el in elements:
        y_mid = (el['y0'] + el['y1']) / 2
        if current_y is None:
            current_y = y_mid
            current_line.append(el)
        else:
            if abs(y_mid - current_y) <= y_tolerance:
                current_line.append(el)
            else:
                lines.append(sorted(current_line, key=lambda e: e['x0']))
                current_line = [el]
                current_y = y_mid
                
    if current_line:
        lines.append(sorted(current_line, key=lambda e: e['x0']))

    return lines

def get_line_bbox(line: List[Dict]) -> Tuple[float, float, float, float]:
    x0 = min([e["x0"] for e in line])
    y0 = min([e["y0"] for e in line])
    x1 = max([e["x1"] for e in line])
    y1 = max([e["y1"] for e in line])
    return (x0, y0, x1, y1)

def get_line_confidence(line: List[Dict]) -> float:
    confs = [e.get("confidence", 1.0) for e in line]
    return sum(confs) / len(confs) if confs else 1.0

def get_element_bbox(elements: List[Dict]) -> Tuple[float, float, float, float]:
    x0 = min([e["x0"] for e in elements])
    y0 = min([e["y0"] for e in elements])
    x1 = max([e["x1"] for e in elements])
    y1 = max([e["y1"] for e in elements])
    return (x0, y0, x1, y1)

# ---------------------------------------------------------------------------
# PARSING
# ---------------------------------------------------------------------------

def is_kv_pair(line: List[Dict]) -> bool:
    line_text = " ".join([e["text"] for e in line]).strip()
    return bool(re.match(r'^[^:]+:\s*.*$', line_text)) or bool(re.match(r'^[^:]+\-\s*.*$', line_text))

def extract_kv_pair(line: List[Dict]) -> Tuple[str, Dict]:
    line_text = " ".join([e["text"] for e in line]).strip()
    m = re.match(r'^([^:]+):\s*(.*)$', line_text)
    if not m:
        m = re.match(r'^([^:]+)\-\s*(.*)$', line_text)
    
    if m:
        k = m.group(1).strip()
        v_text = m.group(2).strip()
        v_elements = []
        for e in line:
            if e["text"] in v_text:
                v_elements.append(e)
        if not v_elements:
            v_elements = line
            
        return k, {
            "value": v_text,
            "confidence": get_line_confidence(v_elements),
            "bbox": get_element_bbox(v_elements)
        }
    return "", {}

def build_rich_cell(elements: List[Dict], default_val: str = "") -> Dict:
    if not elements:
        return {"value": default_val, "confidence": 1.0, "bbox": None}
    val = " ".join([e["text"] for e in elements]).strip()
    return {
        "value": val or default_val,
        "confidence": get_line_confidence(elements),
        "bbox": get_element_bbox(elements)
    }

def process_document(elements: List[Dict]) -> Dict:
    if not elements:
        return {"success": True, "document": {"type": "General", "isStructured": False, "confidence": 0.0, "details": {}}, "sections": []}

    elements = merge_boxes(elements)
    lines = group_into_lines(elements)
    
    # 1. Block Segmentation
    blocks = []
    current_block = []
    last_y = None
    
    for line in lines:
        _, y0, _, y1 = get_line_bbox(line)
        y_mid = (y0 + y1) / 2
        
        if last_y is None:
            current_block.append(line)
        else:
            gap = y0 - last_y
            # Start new block if gap is large (> 30px)
            if gap > 30:
                blocks.append(current_block)
                current_block = [line]
            else:
                current_block.append(line)
        last_y = y1
        
    if current_block:
        blocks.append(current_block)
        
    sections = []
    global_details = {}
    
    table_count = 1
    kv_count = 1
    text_count = 1
    
    # 2. Block Classification and Extraction
    for block in blocks:
        # Check if table (strict heuristic: multiple rows with >= 3 aligned columns)
        counts = [len(l) for l in block]
        max_cols = max(counts) if counts else 0
        
        is_table = False
        if max_cols >= 3:
            rows_with_cols = sum(1 for c in counts if c >= 3)
            if rows_with_cols >= 2:
                is_table = True
        elif max_cols == 2:
            # 2 columns could be KV pairs, but if they have table-like headers and many rows, allow it
            rows_with_cols = sum(1 for c in counts if c == 2)
            if rows_with_cols >= 3:
                is_table = True
                
        if is_table:
            # Simple column clustering based on x-coordinates of the widest line
            widest_line = max(block, key=len)
            col_centers = [(e["x0"] + e["x1"]) / 2 for e in widest_line]  # type: ignore
            col_centers.sort()
            
            headers = [f"Col {i+1}" for i in range(len(col_centers))]
            
            # If the first row looks like headers (words, no numbers)
            first_line_text = " ".join([e["text"] for e in block[0]])
            if not re.search(r'\d', first_line_text):
                # Map first line to columns
                for el in block[0]:
                    c = (el["x0"] + el["x1"]) / 2
                    closest_idx = min(range(len(col_centers)), key=lambda i: abs(col_centers[i] - c))
                    headers[closest_idx] = el["text"]
            
            rows = []
            for i, line in enumerate(block):
                # Skip header row if we used it
                if i == 0 and not re.search(r'\d', first_line_text):
                    continue
                    
                row_data = {h: {"value": "", "confidence": 1.0, "bbox": None} for h in headers}
                row_elements = {h: [] for h in headers}
                
                for el in line:
                    c = (el["x0"] + el["x1"]) / 2
                    closest_idx = min(range(len(col_centers)), key=lambda j: abs(col_centers[j] - c))
                    # Allow slight misalignment
                    if abs(col_centers[closest_idx] - c) < 200:
                        h = headers[closest_idx]
                        row_elements[h].append(el)
                
                for h in headers:
                    row_data[h] = build_rich_cell(row_elements[h])
                    
                rows.append([row_data[h] for h in headers])
            
            sections.append({
                "title": f"Table {table_count}",
                "type": "table",
                "headers": headers,
                "rows": rows
            })
            table_count += 1
            continue
            
        # Check if KV pairs
        kv_pairs = {}
        kv_rows = []
        is_block_kv = False
        
        for line in block:
            if is_kv_pair(line):
                k, v_obj = extract_kv_pair(line)
                if k and v_obj.get("value"):
                    kv_pairs[k] = v_obj
                    is_block_kv = True
        
        if is_block_kv:
            for k, v in kv_pairs.items():
                kv_rows.append([
                    {"value": k, "confidence": 1.0, "bbox": None},
                    v
                ])
                global_details[k] = v["value"]
            
            sections.append({
                "title": f"Document Details {kv_count}" if kv_count > 1 else "Document Details",
                "type": "table",
                "headers": ["Field", "Value"],
                "rows": kv_rows
            })
            kv_count += 1
            continue
            
        # Otherwise Text Paragraph
        text_lines = []
        for line in block:
            text_lines.append(build_rich_cell(line))
            
        sections.append({
            "title": f"Text Block {text_count}",
            "type": "table", # Rendered as a single column table for editability
            "headers": ["Content"],
            "rows": [[l] for l in text_lines]
        })
        text_count += 1

    avg_conf = get_line_confidence(elements)

    res = {
        "success": True,
        "document": {
            "type": "General",
            "module": "General",
            "isRelated": False,
            "isStructured": len(sections) > 0,
            "tableDetected": table_count > 1,
            "confidence": round(avg_conf, 4),
            "pageCount": 1,
            "details": global_details
        },
        "sections": sections,
        "tables": sections,
        "rawText": "\n".join([" ".join([e["text"] for e in l]) for l in lines])
    }
    
    return res

def process_docx_document(file_path: str) -> Dict:
    import docx as _docx
    document = _docx.Document(file_path)
    sections: List[Dict] = []
    kv_pairs: Dict[str, str] = {}
    raw_lines = []
    
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        raw_lines.append(text)
        _KV_COLON = re.compile(r'^(?P<key>[^:]+):\s*(?P<val>.*)$')
        m = _KV_COLON.match(text)
        if m:
            key = m.group("key").strip()
            val = m.group("val").strip()
            if len(key) >= 2 and len(val) >= 1:
                kv_pairs[key] = val
                
    for tbl_idx, table in enumerate(document.tables):
        if not table.rows:
            continue
        header_row = table.rows[0]
        columns = [cell.text.strip() for cell in header_row.cells if cell.text.strip()]
        if not columns:
            continue
        rows = []
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = [cell.text.strip() for cell in row.cells]
            if not any(cells):
                continue
            row_data = []
            for col_idx in range(len(columns)):
                cell_val = cells[col_idx] if col_idx < len(cells) else ""
                row_data.append({"value": cell_val, "confidence": 1.0, "bbox": None})
            rows.append(row_data)
            
        if rows:
            sections.append({
                "title": f"Table {tbl_idx + 1}",
                "type": "table",
                "headers": columns,
                "rows": rows
            })
            
    if kv_pairs:
        sections.insert(0, {
            "title": "Document Details",
            "type": "table",
            "headers": ["Field", "Value"],
            "rows": [[{"value": k, "confidence": 1.0, "bbox": None}, {"value": v, "confidence": 1.0, "bbox": None}] for k, v in kv_pairs.items()]
        })
        
    if not sections and raw_lines:
        sections.append({
            "title": "Extracted Content",
            "type": "table",
            "headers": ["Content"],
            "rows": [[{"value": line, "confidence": 1.0, "bbox": None}] for line in raw_lines if line.strip()]
        })
        
    return {
        "success": True,
        "document": {
            "type": "DOCX",
            "module": "General",
            "isRelated": False,
            "isStructured": len(sections) > 0,
            "tableDetected": len(sections) > 0,
            "confidence": 1.0,
            "pageCount": 1,
            "details": kv_pairs
        },
        "sections": sections,
        "tables": sections,
        "rawText": "\n".join(raw_lines)
    }
